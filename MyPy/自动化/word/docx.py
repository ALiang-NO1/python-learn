from docx import Document

doc = Document('test.docx')
# print(doc.paragraphs)   # 段落地址列表
# for paragraph in doc.paragraphs:
#     print(paragraph.test)   # 打印所有段落的内容

# print("--------提取一些段落内容---------")
# p = doc.paragraphs[0]
# runs = p.runs
# print("runs:", runs)
# for run in runs:
#     print(run.test)

# print("----------添加内容----------")
# h = doc.add_heading("添加的内容", level=1)
# p1 = doc.add_paragraph("添加的第一段内容")
# p2 = doc.add_paragraph("添加的第二段内容")
# doc.save("test.docx")
# print("----------先添加空段，再处理----------")
# p3 = doc.add_paragraph()
# p3.add_run("正常内容！")
# p3.add_run("我是被加粗的内容！").bold = True
# p3.add_run("我是被斜体的内容！").italic = True
# doc.save("text.docx")

# print("-----------添加一张空页, 图片-----------")
# doc.add_page_break()
# doc.add_picture("1.jpg")
# doc.save("test.docx")

# print("---------添加表格----------")
# list1 = [["姓名", "性别", "家庭地址"],
#             ["唐僧", "男", "湖北省"],
#             ["孙悟空", "男", "北京市"],
#             ["猪八戒", "男", "广东省"],
#             ["沙和尚", "男", "湖南省"]]
# list2 = [["姓名", "性别", "家庭地址"],
#             ["貂蝉", "女", "河北省"],
#             ["杨贵妃", "女", "贵州省"],
#             ["西施", "女", "山东省"]]
# table1 = doc.add_table(rows=5, cols=3)
# for row in range(5):
#     cells = table1.rows[row].cells
#     for col in range(3):
#         cells[col].text = str(list1[row][col])

# table2 = doc.add_table(rows=5, cols=3)
# for row in range(4):
#     cells = table2.rows[row].cells
#     for col in range(3):
#         cells[col].text = str(list1[row][col])
# doc.save("test.docx")

print("---------将内容保存到excel中---------")
from openpyxl import Workbook
t0 = doc.tables[0]
workbook = Workbook()
sheet = workbook.active
for i in range(len(t0.rows)):
    list1 = []
    for j in range(len(t0.columns)):
        list1.append(t0.cell(i, j).text)
        sheet.append(list1)
workbook.save("E:\python文档\my_py\自动化\excel\\test.xlsx")
print("复制表格成功！")


# # doc.save("test.docx")
# for paragraph in doc.paragraphs:
#     print(paragraph.text)   # 打印所有段落的内容